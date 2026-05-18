"""
Omysha Foundation — Candidate Evaluation Report Generator (Word Document)

Generates a formatted .docx report from structured evaluation data passed as JSON.

Usage:
  python generate-docx.py <input.json> <output.docx>

Input JSON structure:
{
  "batch_info": {
    "interview_date": "March 16, 2026",
    "interviewers": "Sushma (Lead), Suhani (Support)",
    "role": "HR Intern",
    "transcript_note": "Optional note about transcript quality"
  },
  "candidates": [
    {
      "name": "Candidate Name",
      "role_applied": "HR Intern",
      "interview_date": "March 16, 2026",
      "availability": "3 PM onwards",
      "stipend_discussed": "Not discussed",
      "resume_provided": "Yes",
      "rounds_completed": "GD + HR Round + Fit Round",
      "confidence_level": "High",
      "limitation_note": "",
      "background": "Education and experience summary...",
      "interview_summary": "4-5 sentence summary...",
      "natural_fit": {
        "role_reference": "natural-fit-hr.md",
        "traits": [
          {
            "trait": "Empathy and Patience",
            "score": 4,
            "observation": "HR observation text...",
            "example": "Behavioral example..."
          }
        ],
        "average": 3.8
      },
      "org_fit": {
        "values": [
          {
            "value": "Ownership and Alignment",
            "score": 4,
            "observation": "HR observation text...",
            "example": "Behavioral example..."
          }
        ],
        "average": 3.8
      },
      "total_score": 7.6,
      "recommendation": "Can Be Considered",
      "recommendation_emoji": "🟢",
      "key_strengths": ["Strength 1", "Strength 2"],
      "areas_for_development": ["Area 1", "Area 2"],
      "overall_observation": "3-5 sentence holistic observation...",
      "recommendation_basis": "Cite specific traits and scores..."
    }
  ],
  "gd_eliminations": [
    {
      "name": "Candidate Name",
      "outcome": "Eliminated after GD",
      "note": "Brief reason..."
    }
  ]
}
"""

import json
import sys
import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shading_elem = tc_pr.makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): color_hex,
            qn("w:val"): "clear",
        },
    )
    tc_pr.append(shading_elem)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2E4057")

    # Data rows
    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, "F0F4F8")

    # Apply column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    return table


def get_score_color(score):
    """Return color based on recommendation."""
    if score >= 8.5:
        return RGBColor(0x27, 0xAE, 0x60)  # Green
    elif score >= 7.0:
        return RGBColor(0x2E, 0x86, 0xC1)  # Blue
    elif score >= 5.5:
        return RGBColor(0xF3, 0x9C, 0x12)  # Yellow/Orange
    elif score >= 3.5:
        return RGBColor(0xE7, 0x4C, 0x3C)  # Red
    else:
        return RGBColor(0x8B, 0x00, 0x00)  # Dark Red


def add_candidate_report(doc, candidate, index):
    """Add a single candidate's evaluation report to the document."""
    name = candidate["name"]

    # Candidate header
    doc.add_page_break()
    heading = doc.add_heading(f"Report {index} — {name}", level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    # --- Section 1: Candidate Details ---
    doc.add_heading("1. Candidate Details", level=2)
    details = [
        ["Candidate Name", candidate.get("name", "—")],
        ["Role Applied", candidate.get("role_applied", "—")],
        ["Interview Date", candidate.get("interview_date", "—")],
        ["Availability Declared", candidate.get("availability", "—")],
        ["Stipend Discussed", candidate.get("stipend_discussed", "—")],
        ["Resume Provided", candidate.get("resume_provided", "—")],
    ]
    add_styled_table(doc, ["Field", "Details"], details, col_widths=[5, 12])

    # --- Section 2: Assessment Confidence ---
    doc.add_heading("2. Assessment Confidence", level=2)
    confidence = candidate.get("confidence_level", "Medium")
    conf_rows = [
        ["Rounds Completed", candidate.get("rounds_completed", "—")],
        ["Confidence Level", confidence],
        ["Limitation Note", candidate.get("limitation_note", "—")],
    ]
    add_styled_table(doc, ["Factor", "Detail"], conf_rows, col_widths=[5, 12])

    if confidence in ("Medium", "Low"):
        warn = doc.add_paragraph()
        run = warn.add_run(
            f"⚠️ Confidence is {confidence} — scores are based on limited evidence."
        )
        run.font.color.rgb = RGBColor(0xF3, 0x9C, 0x12)
        run.bold = True

    # --- Section 3: Background ---
    doc.add_heading("3. Candidate Background", level=2)
    doc.add_paragraph(candidate.get("background", "Not available."))

    # --- Section 4: Interview Summary ---
    doc.add_heading("4. Interview Summary", level=2)
    doc.add_paragraph(candidate.get("interview_summary", "Not available."))

    # --- Section 5: Natural Fit ---
    nf = candidate.get("natural_fit", {})
    nf_title = "5. Natural Fit Evaluation"
    if confidence in ("Medium", "Low"):
        nf_title += f"  ⚠️ Limited evidence — confidence: {confidence}"
    doc.add_heading(nf_title, level=2)

    ref_para = doc.add_paragraph()
    ref_para.add_run("Role reference used: ").bold = True
    ref_para.add_run(nf.get("role_reference", "—"))

    traits = nf.get("traits", [])
    trait_rows = []
    for t in traits:
        trait_rows.append([
            t.get("trait", "—"),
            f"{t.get('score', '—')}/5",
            t.get("observation", "—"),
        ])
    add_styled_table(
        doc,
        ["Trait", "Score", "Evaluation"],
        trait_rows,
        col_widths=[4, 1.5, 11.5],
    )

    nf_avg = nf.get("average", 0)
    avg_para = doc.add_paragraph()
    run = avg_para.add_run(f"Natural Fit Average: {nf_avg} / 5")
    run.bold = True
    run.font.size = Pt(11)

    # --- Section 6: Org Fit ---
    of = candidate.get("org_fit", {})
    of_title = "6. Organizational Fit Evaluation"
    if confidence in ("Medium", "Low"):
        of_title += f"  ⚠️ Limited evidence — confidence: {confidence}"
    doc.add_heading(of_title, level=2)

    values = of.get("values", [])
    value_rows = []
    for v in values:
        value_rows.append([
            v.get("value", "—"),
            f"{v.get('score', '—')}/5",
            v.get("observation", "—"),
        ])
    add_styled_table(
        doc,
        ["Value", "Score", "Evaluation"],
        value_rows,
        col_widths=[4, 1.5, 11.5],
    )

    of_avg = of.get("average", 0)
    avg_para = doc.add_paragraph()
    run = avg_para.add_run(f"Org Fit Average: {of_avg} / 5")
    run.bold = True
    run.font.size = Pt(11)

    # --- Section 7: Score Summary ---
    doc.add_heading("7. Score Summary", level=2)
    total = candidate.get("total_score", 0)
    score_rows = [
        ["Natural Fit", f"{nf_avg} / 5", "50%", f"{nf_avg}"],
        ["Org Fit", f"{of_avg} / 5", "50%", f"{of_avg}"],
        ["TOTAL SCORE", "", "", f"{total} / 10"],
    ]
    add_styled_table(
        doc,
        ["Component", "Raw Score", "Weight", "Contribution"],
        score_rows,
        col_widths=[4, 3, 3, 3],
    )

    rec = candidate.get("recommendation", "—")
    emoji = candidate.get("recommendation_emoji", "")
    rec_para = doc.add_paragraph()
    run = rec_para.add_run(f"Score-based recommendation: {emoji} {rec}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = get_score_color(total)

    # --- Section 8: Key Strengths ---
    doc.add_heading("8. Key Strengths", level=2)
    for s in candidate.get("key_strengths", []):
        doc.add_paragraph(s, style="List Bullet")

    # --- Section 9: Areas for Development ---
    doc.add_heading("9. Areas for Development", level=2)
    for a in candidate.get("areas_for_development", []):
        doc.add_paragraph(a, style="List Bullet")

    # --- Section 10: Overall HR Observation ---
    doc.add_heading("10. Overall HR Observation", level=2)
    doc.add_paragraph(candidate.get("overall_observation", "Not available."))

    # --- Section 11: Recommendation ---
    doc.add_heading("11. Recommendation", level=2)

    # Action label is the PRIMARY decision signal — shown first, largest
    action_label = candidate.get("action_label", "")
    if action_label:
        action_para = doc.add_paragraph()
        action_run = action_para.add_run(action_label)
        action_run.bold = True
        action_run.font.size = Pt(16)
        action_run.font.color.rgb = get_score_color(total)
    else:
        # Fallback: show score tier if no action label set
        rec_box = doc.add_paragraph()
        run = rec_box.add_run(f"{emoji} {rec}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = get_score_color(total)

    # Score and tier shown as supporting context, smaller
    score_line = doc.add_paragraph()
    score_line.add_run("Score: ").bold = True
    score_run = score_line.add_run(f"{total} / 10  |  {rec}")
    score_run.font.size = Pt(10)

    # Key Risk
    key_risk = candidate.get("key_risk", "")
    if key_risk:
        risk_para = doc.add_paragraph()
        risk_para.add_run("Key Risk: ").bold = True
        risk_run = risk_para.add_run(key_risk)
        risk_run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)

    # Placement Fit (only for Good/Strong candidates)
    placement_fit = candidate.get("placement_fit", "")
    if placement_fit:
        fit_para = doc.add_paragraph()
        fit_para.add_run("Placement Fit: ").bold = True
        fit_para.add_run(placement_fit)

    basis_line = doc.add_paragraph()
    basis_line.add_run("Basis: ").bold = True
    basis_line.add_run(candidate.get("recommendation_basis", "—"))


def generate_report(data, output_path):
    """Generate the full Word document from evaluation data."""
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # --- Title Page ---
    doc.add_paragraph()  # spacing
    title = doc.add_heading("Candidate Evaluation Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Omysha Foundation — Confidential")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    doc.add_paragraph()  # spacing

    # Batch info
    batch = data.get("batch_info", {})
    info_rows = [
        ["Interview Date", batch.get("interview_date", "—")],
        ["Interviewer(s)", batch.get("interviewers", "—")],
        ["Role", batch.get("role", "—")],
    ]
    add_styled_table(doc, ["Field", "Details"], info_rows, col_widths=[5, 12])

    # Transcript note
    tn = batch.get("transcript_note", "")
    if tn:
        doc.add_paragraph()
        note_para = doc.add_paragraph()
        run = note_para.add_run(f"Transcript Note: {tn}")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    # --- GD Eliminations (if any) ---
    gd_elims = data.get("gd_eliminations", [])
    if gd_elims:
        doc.add_page_break()
        doc.add_heading("Candidates Eliminated at GD / Pre-Interview", level=1)
        doc.add_paragraph(
            "The following candidates did not proceed past Round 1 and have "
            "insufficient data for full evaluation."
        )
        for elim in gd_elims:
            doc.add_heading(
                f"{elim.get('name', '—')} — {elim.get('outcome', 'Eliminated')}",
                level=3,
            )
            doc.add_paragraph(elim.get("note", "No details available."))

    # --- Individual Candidate Reports ---
    candidates = data.get("candidates", [])
    for i, candidate in enumerate(candidates, 1):
        add_candidate_report(doc, candidate, i)

    # --- Comparative Summary (if multiple candidates) ---
    if len(candidates) > 1:
        doc.add_page_break()
        doc.add_heading("Comparative Summary", level=1)

        summary_rows = []
        sorted_candidates = sorted(
            candidates, key=lambda c: c.get("total_score", 0), reverse=True
        )
        for rank, c in enumerate(sorted_candidates, 1):
            summary_rows.append([
                str(rank),
                c.get("name", "—"),
                c.get("role_applied", "—"),
                f"{c.get('natural_fit', {}).get('average', '—')} / 5",
                f"{c.get('org_fit', {}).get('average', '—')} / 5",
                f"{c.get('total_score', '—')} / 10",
                f"{c.get('recommendation_emoji', '')} {c.get('recommendation', '—')}",
            ])
        add_styled_table(
            doc,
            ["Rank", "Name", "Role", "NF Avg", "OF Avg", "Total", "Recommendation"],
            summary_rows,
            col_widths=[1.5, 3, 2.5, 2, 2, 2, 3.5],
        )

        # Flag close scores
        for i in range(len(sorted_candidates) - 1):
            s1 = sorted_candidates[i].get("total_score", 0)
            s2 = sorted_candidates[i + 1].get("total_score", 0)
            if abs(s1 - s2) <= 0.3:
                warn = doc.add_paragraph()
                run = warn.add_run(
                    f"⚠️ {sorted_candidates[i]['name']} and "
                    f"{sorted_candidates[i+1]['name']} are within 0.3 points — "
                    f"recommend HR OL review both."
                )
                run.font.color.rgb = RGBColor(0xF3, 0x9C, 0x12)
                run.bold = True

        # GD eliminations in summary
        if gd_elims:
            doc.add_paragraph()
            doc.add_heading("GD-Only Candidates", level=2)
            gd_rows = []
            for e in gd_elims:
                gd_rows.append([
                    e.get("name", "—"),
                    e.get("outcome", "—"),
                    e.get("note", "—"),
                ])
            add_styled_table(
                doc,
                ["Name", "Outcome", "Note"],
                gd_rows,
                col_widths=[4, 4, 9],
            )

    # --- Footer on last page ---
    doc.add_paragraph()
    doc.add_paragraph("—" * 60)
    footer_lines = [
        "Final decision authority: HR OL — Nitin Sir",
        "This report is a scoring input, not a hiring decision.",
        "To check rejection criteria (H1–H5, S1–S7), use /justify-rejection.",
        "Omysha Foundation — Confidential",
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
    ]
    for line in footer_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    # Save
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python generate-docx.py <input.json> <output.docx>")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    with open(input_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    result = generate_report(data, output_path)
    print(f"Report generated: {result}")
