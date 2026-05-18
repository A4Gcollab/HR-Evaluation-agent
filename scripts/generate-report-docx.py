"""
Generate HR Intern Evaluation Report — Word Document
Batch: March 2026 (Transcript provided by user)
Organization: Omysha Foundation

Round-gating applied: Only scores NF/OF when those rounds were actually conducted.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)


def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
    return table


def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')


def add_italic_note(text, size=9):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(size)
    return p


def calc_score(nf_scores, of_scores):
    nf_avg = round(sum(nf_scores) / len(nf_scores), 1)
    of_avg = round(sum(of_scores) / len(of_scores), 1)
    total = round(((nf_avg + of_avg) / 2) * 2, 1)
    if total >= 8.5:
        rec = "Recommended"
    elif total >= 7.0:
        rec = "Can Be Considered"
    elif total >= 5.5:
        rec = "Borderline"
    elif total >= 3.5:
        rec = "Not Recommended"
    else:
        rec = "Strong No"
    return nf_avg, of_avg, total, rec


# ============================================================
# COVER PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('HR Intern Evaluation Report', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Omysha Foundation — Confidential')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Batch Interview Date: ').bold = True
info.add_run('March 2026\n')
info.add_run('Report Generated: ').bold = True
info.add_run('March 23, 2026\n')
info.add_run('Interviewer(s): ').bold = True
info.add_run('Sushma (Lead HR), Suhani (Natural Fit), Nausheen (Org Fit)\n')
info.add_run('Role: ').bold = True
info.add_run('HR Intern\n')
info.add_run('Total Participants: ').bold = True
info.add_run('8')

doc.add_paragraph()
add_italic_note(
    'Transcript Note: This transcript is a Zoom AI Companion auto-generated VTT file. '
    'Transcription errors are present throughout (e.g., "Omisha" = Omysha, "Wong" = VONG, '
    '"A4G/FOG/F4G" = AI4Good). Evaluations account for transcription artifacts and focus '
    'on behavioral signals rather than exact wording.', 9
)

doc.add_page_break()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
add_heading_styled('Executive Summary', level=1)

doc.add_paragraph(
    'This report evaluates 8 candidates who participated in the HR Intern group interview '
    'conducted via Zoom. The interview consisted of up to 3 rounds: Group Discussion (Round 1), '
    'HR Round (Round 2), and Natural Fit + Organizational Fit assessment (Round 3).'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Scoring is only applied to rounds that were actually conducted. ')
run.bold = True
p.add_run(
    'Candidates who did not complete the Natural Fit and Organizational Fit rounds '
    'receive qualitative assessments only — no NF/OF scores and no Total Score.'
)

doc.add_paragraph()
add_table(
    ["Candidate", "Rounds Completed", "Scored?", "Total Score", "Recommendation"],
    [
        ["Vidhi Upadhyay", "GD + HR + NF + OF", "Yes", "7.2 / 10", "Can Be Considered"],
        ["Drishti", "GD + HR (partial)", "No", "—", "Qualitative only"],
        ["Divya Patidar", "GD (advanced to next round*)", "No", "—", "Qualitative only"],
        ["Ansh Prakash", "GD only (eliminated)", "No", "—", "GD summary only"],
        ["Suraksha Chakraborty", "GD only (eliminated)", "No", "—", "GD summary only"],
        ["Mansi", "GD only (left)", "No", "—", "GD summary only"],
        ["Ritul Goyal", "None (joined late)", "No", "—", "Not evaluated"],
        ["Bhavaniprasad Gedewad", "None (technical)", "No", "—", "Not evaluated"],
    ]
)

doc.add_paragraph()
add_italic_note(
    '* Divya Patidar advanced to the next round per interviewer decision, but that round '
    'is not captured in this transcript. Score would require that data.'
)

doc.add_page_break()

# ============================================================
# NOT EVALUATED
# ============================================================
add_heading_styled('Candidates Not Evaluated', level=1)

add_heading_styled('Ritul Goyal — Joined Late, No Participation', level=3)
doc.add_paragraph(
    'Joined the meeting late after the Group Discussion had already concluded. '
    'Was sent to a breakout room. Did not participate in any evaluation round. '
    'No evaluation possible.'
)

add_heading_styled('Bhavaniprasad Gedewad — Technical Issues, No Participation', level=3)
doc.add_paragraph(
    'Appeared briefly with audio/technical issues at the start of the meeting. '
    'Did not participate in the Group Discussion or any subsequent round. '
    'No evaluation possible.'
)

doc.add_page_break()

# ============================================================
# GD-ONLY ELIMINATIONS (no scores — summaries only)
# ============================================================
add_heading_styled('GD-Only Candidates — Summary Entries', level=1)

doc.add_paragraph(
    'The following candidates participated in Round 1 (Group Discussion) only. '
    'They were either eliminated after GD or left before the HR Round. '
    'Since no Natural Fit or Organizational Fit round was conducted, no NF/OF '
    'scores are assigned. These entries contain GD behavioral observations only.'
)

# --- ANSH ---
doc.add_paragraph()
add_heading_styled('Ansh Prakash — Eliminated After GD', level=2)
add_table(
    ["Field", "Details"],
    [
        ["Rounds Completed", "GD (Round 1) only"],
        ["Outcome", "Eliminated after Group Discussion"],
    ]
)
doc.add_paragraph()
add_heading_styled('GD Observations', level=3)
doc.add_paragraph(
    'Ansh was the most proactive GD participant. He volunteered to speak first on AI4Good '
    'when Divya hesitated, and provided the most detailed contributions in the batch. He '
    'identified the three strategic focus areas of AI4Good (AI Swaraj model, AI for Mental '
    'Well-being, Sustainability/Socioeconomic growth), described VONG as a collaborative '
    'youth movement, and discussed Wongsters\' focus on education, advocacy, and action. '
    'He was corrected by Sushma on terminology ("Sankals" vs "Wongsters") and adapted '
    'immediately and gracefully.'
)
add_heading_styled('Behavioral Signals', level=3)
add_bullet('Communication: Strong — organized points into clear structures, filled silences proactively')
add_bullet('Confidence: Strong — volunteered multiple times, spoke with conviction')
add_bullet('Org Awareness: Strong — most researched candidate, knew specific focus areas')
add_bullet('Attitude: Positive — accepted correction openly, collaborative tone')
add_italic_note('Natural Fit / Org Fit: Round not conducted — not scored.')

# --- SURAKSHA ---
doc.add_paragraph()
doc.add_page_break()
add_heading_styled('Suraksha Chakraborty — Eliminated After GD', level=2)
add_table(
    ["Field", "Details"],
    [
        ["Rounds Completed", "GD (Round 1) only"],
        ["Outcome", "Eliminated after Group Discussion"],
    ]
)
doc.add_paragraph()
add_heading_styled('GD Observations', level=3)
doc.add_paragraph(
    'Suraksha delivered some of the most insightful GD contributions in the batch. She '
    'framed AI4Good as breaking the myth that AI replaces jobs, connected the three pillars '
    'to a chain (governance = nation, sustainability = society, well-being = self), and '
    'described HR as "a leader who is not in front, but in back, setting the whole table." '
    'She also connected remote work to youth empowerment and discussed AI simplifying '
    'complex governance and laws for common people — not just experts.'
)
add_heading_styled('Behavioral Signals', level=3)
add_bullet('Communication: Strong — clear, confident delivery, structured points well')
add_bullet('Confidence: Strong — spoke multiple times without hesitation')
add_bullet('Org Awareness: Strong — connected deeply with mission, original chain metaphor')
add_bullet('Attitude: Positive — collaborative, acknowledged and built on peers by name')
add_italic_note('Natural Fit / Org Fit: Round not conducted — not scored.')

# --- MANSI ---
doc.add_paragraph()
add_heading_styled('Mansi — Left Before HR Round', level=2)
add_table(
    ["Field", "Details"],
    [
        ["Rounds Completed", "GD (Round 1) only"],
        ["Outcome", "Left the interview before being called for further rounds"],
    ]
)
doc.add_paragraph()
add_heading_styled('GD Observations', level=3)
doc.add_paragraph(
    'Mansi was an active GD participant who demonstrated solid understanding of the '
    'organization. She mentioned the four pillars of AI4Good, gave a practical example '
    'of AI in HR recruitment (resume screening reducing recruiter workload), discussed '
    'L&D initiatives for employees, and connected youth empowerment to the nation\'s '
    'future. Her contributions were relevant and showed HR domain familiarity.'
)
add_heading_styled('Behavioral Signals', level=3)
add_bullet('Communication: Moderate — clear and organized, sometimes verbose')
add_bullet('Confidence: Moderate — participated actively but did not lead')
add_bullet('Org Awareness: Moderate — understood mission, practical HR-AI connection')
add_bullet('Attitude: Positive — polite sign-off, collaborative throughout')
add_italic_note('Natural Fit / Org Fit: Round not conducted — not scored.')

# --- DIVYA ---
doc.add_paragraph()
doc.add_page_break()
add_heading_styled('Divya Patidar — GD Only (Advanced to Next Round)', level=2)
add_table(
    ["Field", "Details"],
    [
        ["Rounds Completed", "GD (Round 1) only in this transcript"],
        ["Outcome", "Advanced to next round — data not captured in this transcript"],
    ]
)
doc.add_paragraph()
add_heading_styled('GD Observations', level=3)
doc.add_paragraph(
    'Divya volunteered to speak first in the GD but was not sufficiently prepared on '
    'AI4Good specifically — she began discussing Omysha Foundation generally and was '
    'redirected by Sushma to focus on FOG. She requested additional preparation time. '
    'Later, she provided a well-structured answer about the HR operations role (onboarding, '
    'documentation, HRMS, employee engagement, performance management, cross-department '
    'coordination). During the break, she was helpful to Vidhi regarding Zoom navigation.'
)
add_heading_styled('Behavioral Signals', level=3)
add_bullet('Communication: Moderate — HR role answer was structured; initial FOG attempt was underprepared')
add_bullet('Confidence: Moderate — volunteered first but couldn\'t deliver initially, recovered later')
add_bullet('Org Awareness: Weak-Moderate — knew about Omysha broadly, not FOG specifically')
add_bullet('Attitude: Positive — accepted redirection, helpful to peers')
add_italic_note(
    'Natural Fit / Org Fit: Round not conducted — not scored. '
    'Divya advanced to the next round per interviewer decision. '
    'A full evaluation requires the transcript of that round.'
)

doc.add_page_break()

# ============================================================
# PARTIAL REPORT — DRISHTI (GD + HR Round, no NF/OF)
# ============================================================
add_heading_styled('REPORT — Drishti (Partial — GD + HR Round Only)', level=1)

add_heading_styled('1. Candidate Details', level=2)
add_table(
    ["Field", "Details"],
    [
        ["Candidate Name", "Drishti"],
        ["Position Applied", "HR Intern"],
        ["Rounds Completed", "GD (Round 1) + HR Round (Round 2, partial)"],
        ["Interview Date", "March 2026"],
        ["Availability Declared", "College 9:30 AM - 3/4 PM. Can reach home by 4 PM."],
        ["Stipend Discussed", "Not discussed — eliminated before stipend conversation"],
        ["Resume Provided", "Yes — content writing internships, NGO (street school), Psychology UG (2nd year)"],
    ]
)

doc.add_paragraph()
add_heading_styled('2. Assessment Confidence', level=2)
add_table(
    ["Factor", "Detail"],
    [
        ["Rounds Completed", "GD + HR Round (partial). Natural Fit and Org Fit rounds were NOT conducted."],
        ["Confidence Level", "LOW"],
        ["Limitation Note", "Eliminated during HR Round due to availability conflict. NF and OF rounds were never administered. No NF/OF scores can be assigned."],
    ]
)

add_heading_styled('3. Candidate Background', level=2)
doc.add_paragraph(
    'Undergraduate psychology student (2nd year), aligned with organizational psychology. '
    'Previous experience in content writing internships and an NGO (street school) — ending '
    'within 2 weeks. Self-described introvert who sees HR as a way to connect with people. '
    'No prior HR-specific experience.'
)

add_heading_styled('4. Interview Summary', level=2)
doc.add_paragraph(
    'Drishti was among the strongest GD contributors — articulate, structured, and collaborative. '
    'She connected AI4Good to collaboration with institutions and policymakers, discussed VONG as '
    'an amplifier for youth voice, and linked the HR role to creating a positive work environment '
    'with values like teamwork and communication. During the HR Round, she presented as self-aware '
    'and genuine — openly discussing her introversion, her interest in organizational psychology, '
    'and her motivation to apply psychology in real situations. The interviewers explicitly noted '
    'they liked her understanding of the organization and valued her psychology background. However, '
    'her college schedule (ending 3-4 PM) created an irreconcilable availability conflict with the '
    'HR role\'s 3 PM requirement.'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Natural Fit Evaluation: Round not conducted — not scored.')
run.bold = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Organizational Fit Evaluation: Round not conducted — not scored.')
run.bold = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Score Summary: Cannot be computed — requires both NF and OF round data.')
run.bold = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

add_heading_styled('5. Key Strengths (from GD + HR Round)', level=2)
add_bullet('Strong GD performance — articulate, structured, collaborative, among the best communicators')
add_bullet('Psychology background aligned with organizational psychology / HR')
add_bullet('Genuine mission alignment — attracted to social impact, not just career advancement')
add_bullet('Self-aware and honest about limitations (introversion, beginner status)')
add_bullet('Interviewers explicitly valued her profile: "we liked your understanding... we\'re looking for psychology background"')

add_heading_styled('6. Areas of Concern', level=2)
add_bullet('Availability conflict: College ends 3-4 PM; HR role requires 3 PM onwards availability')
add_bullet('No prior HR experience — entirely learning-stage')
add_bullet('Content writing background; HR transition is early-stage')

add_heading_styled('7. Overall HR Observation', level=2)
doc.add_paragraph(
    'Drishti demonstrated strong communication and analytical skills in the GD, and genuine '
    'motivation during the HR Round. Her psychology background is a valuable asset for the HR '
    'role, and interviewers noted this positively. The elimination was driven solely by a '
    'scheduling constraint — not by performance or fit concerns. If her availability changes '
    '(e.g., after semester ends or schedule shifts), she would merit a full re-evaluation '
    'including NF and OF rounds.'
)

add_heading_styled('8. Recommendation', level=2)
p = doc.add_paragraph()
run = p.add_run('Eliminated — Availability Conflict')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0xCC, 0x99, 0x00)
doc.add_paragraph(
    'Outcome: Eliminated during HR Round. College schedule (9:30 AM - 3/4 PM) is incompatible '
    'with the HR role\'s 3 PM onwards availability requirement.'
)
doc.add_paragraph(
    'Qualitative assessment: Strong GD performance and genuine HR interest with a relevant '
    'psychology background. Not a performance-based elimination. Re-evaluate if availability '
    'constraint resolves.'
)
add_italic_note('No score-based recommendation possible — NF + OF round data required for scoring.')

doc.add_paragraph()
add_italic_note('Final decision authority: HR OL — Nitin Sir')
add_italic_note('This report is a scoring input, not a hiring decision.')

doc.add_page_break()

# ============================================================
# FULL REPORT — VIDHI UPADHYAY (All rounds conducted)
# ============================================================
add_heading_styled('REPORT — Vidhi Upadhyay (Full Evaluation)', level=1)

add_heading_styled('1. Candidate Details', level=2)
add_table(
    ["Field", "Details"],
    [
        ["Candidate Name", "Vidhi Upadhyay"],
        ["Position Applied", "HR Intern"],
        ["Rounds Completed", "GD (Round 1) + HR Round (Round 2) + Natural Fit (Suhani) + Org Fit (Nausheen)"],
        ["Interview Date", "March 2026"],
        ["Availability Declared", "Fully available. College (B.Com) ends April 10. Exams April 5-10. Confirmed 3 PM - 9/10 PM."],
        ["Stipend Discussed", "Yes — understood structure (5K / 7.5K / 10K based on performance >= 3.5/5 and 30 hrs/week)"],
        ["Resume Provided", "Yes"],
    ]
)

doc.add_paragraph()
add_heading_styled('2. Assessment Confidence', level=2)
add_table(
    ["Factor", "Detail"],
    [
        ["Rounds Completed", "All rounds: GD + HR Round + Natural Fit (Suhani) + Org Fit (Nausheen)"],
        ["Confidence Level", "HIGH"],
        ["Limitation Note", "Full evidence base. Transcript is auto-generated VTT with minor transcription artifacts."],
    ]
)

add_heading_styled('3. Candidate Background', level=2)
doc.add_paragraph(
    'Final-year B.Com student at KJ Somaiya College. Exams end April 10, 2026. Currently '
    'interning at an NGO (cancer patients focus) as both HR and project management — commitment '
    'ends March 23. Has conducted 50-70 interviews and hired 10-15 candidates in 2-3 months. '
    'Initially hired as project manager but transitioned to HR based on communication skills. '
    'Also managed performance evaluation sheets at the NGO. No prior paid experience (NGO was '
    'unpaid/volunteer).'
)

add_heading_styled('4. Interview Summary', level=2)
doc.add_paragraph(
    'Vidhi presented as an enthusiastic, communication-oriented candidate with genuine HR interest '
    'and relevant practical experience. In the GD, she had initial audio issues but recovered and '
    'contributed points about AI efficiency and manpower optimization — competent but less detailed '
    'than top GD performers. The HR Round revealed strong practical grounding: she described her NGO '
    'experience with specificity (50-70 interviews, 10-15 hires, performance sheet management), '
    'showed clear career direction toward HR, and demonstrated full availability with no competing '
    'commitments. She understood the stipend structure (after correction on first-tier amount), '
    'confirmed 6-month commitment and 30 hours/week, and expressed willingness to take on the '
    'secondary Community Building role. During the Natural Fit round, her answers showed empathetic '
    'instincts but lacked structured HR frameworks. The Org Fit round revealed practical '
    'experience-based answers that Nausheen found satisfactory. Internal interviewer discussion '
    'noted she is "good, not very good" — a learning-stage candidate with potential.'
)

add_heading_styled('5. Natural Fit Evaluation', level=2)
add_italic_note('Role: HR Intern | Reference: HR Role Natural Fit Dimensions | Round conducted by: Suhani')

nf_vidhi = [4, 3, 3, 3, 4]
add_table(
    ["Trait", "Score (1-5)", "HR Observation", "Behavioral Example"],
    [
        ["Empathy and Patience", "4",
         "Consistently people-oriented. First instinct for struggling employees is one-on-one connection and understanding root cause before formal action.",
         "On employee crisis: 'Before taking a formal approach, I would try to connect with them one-on-one and understand the reason behind their poor performance.'"],
        ["Communication & Active Listening", "3",
         "Enthusiastic and articulate but sometimes verbose. Needed questions repeated or explained (authenticity, conflict vs meeting). Good rapport-building but could be more precise.",
         "When Nausheen asked about authenticity, candidate asked for clarification — shows willingness to seek clarity but also comprehension gaps."],
        ["Problem Solving & Decision Making", "3",
         "Practical reasoning but lacked structured frameworks. Chose management meeting over urgent conflict with cause-and-effect reasoning. HR noted she was more analytical in person than transcript suggests.",
         "On conflict vs meeting: chose management meeting first reasoning that 'fixing management reduces employee conflicts' — unconventional but logical."],
        ["Initiate & Set Company Culture", "3",
         "Proposed appreciation posts at NGO to boost morale — an implemented initiative, not hypothetical. Willing to take on CB role. Culture-building instinct but limited HR-specific strategic thinking.",
         "Created weekly appreciation posts at NGO, shared on LinkedIn — an implemented idea that improved team morale."],
        ["Discretion & Professionalism", "4",
         "Honest about experience level, didn't overstate capabilities. Committed clearly to 6-month tenure. Handled stipend correction gracefully. Would not break commitment for better offers.",
         "On commitment: 'If I am already selected here, I have to commit within the tenure... I would not go to any other opportunities within these 6 months.'"],
    ]
)
nf_avg = round(sum(nf_vidhi) / 5, 1)
p = doc.add_paragraph()
run = p.add_run(f'Natural Fit Average: {nf_avg} / 5')
run.bold = True

add_heading_styled('6. Organizational Fit Evaluation', level=2)
add_italic_note('Round conducted by: Nausheen')

of_vidhi = [4, 3, 4, 3, 4]
add_table(
    ["Value", "Score (1-5)", "HR Observation", "Behavioral Example"],
    [
        ["Ownership & Alignment", "4",
         "Strong mission alignment — attracted to youth impact and social empowerment. NGO background shows pattern of purpose-driven work. Comfortable working without spoon-feeding.",
         "On motivation: 'Umesha's main motive, youth impact and social empowerment... as a youth, I really want to experience that.' Comfortable with independent execution."],
        ["Respect & Acceptance", "3",
         "Respectful throughout. Accepted stipend correction without defensiveness. Collaborative approach to disagreements. Did not dominate conversation.",
         "On working with someone different: described letting colleague take their part, then sharing own perspective — balanced approach."],
        ["Innovation & Imagination", "4",
         "Proposed and implemented appreciation posts initiative. Adapted from WhatsApp-based to Google Sheet-based systems. Practical creativity. Nausheen noted this positively.",
         "Adapted from unstructured WhatsApp communication to structured Google Sheets tracking: 'I was able to get involved in both systems very naturally.'"],
        ["Agility", "3",
         "Adapted from project management to HR. Handled audio issues in GD. But needed questions explained multiple times — minor comprehension agility gaps.",
         "Career pivot from project management to HR after discovering affinity through interviews. Adapted to different work systems at NGO."],
        ["Integrity & Authenticity", "4",
         "Consistently honest: about unpaid NGO work, early-career stage, exam schedule, needing guidance for bigger decisions. Nausheen was satisfied with authenticity answers.",
         "On mistakes: 'I would definitely take the accountability... I would try to make changes as soon as possible.' Did not fabricate experience."],
    ]
)
of_avg = round(sum(of_vidhi) / 5, 1)
p = doc.add_paragraph()
run = p.add_run(f'Organizational Fit Average: {of_avg} / 5')
run.bold = True

add_heading_styled('7. Score Summary', level=2)
nf_avg_v, of_avg_v, total_v, rec_v = calc_score(nf_vidhi, of_vidhi)

add_table(
    ["Component", "Raw Score", "Weight", "Contribution"],
    [
        ["Natural Fit", f"{nf_avg_v} / 5", "50%", str(nf_avg_v)],
        ["Organizational Fit", f"{of_avg_v} / 5", "50%", str(of_avg_v)],
        ["TOTAL SCORE", "", "", f"{total_v} / 10"],
    ]
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(f'Score-based Recommendation: {rec_v}')
run.bold = True
run.font.size = Pt(11)

add_heading_styled('8. Key Strengths', level=2)
add_bullet('Genuine HR interest with relevant practical experience (50-70 interviews, 10-15 hires)')
add_bullet('Strong empathetic instinct — first approach is always to understand before acting (NF: Empathy 4/5)')
add_bullet('Full availability confirmed — no schedule conflicts, 6-month commitment, 30 hrs/week')
add_bullet('Mission-aligned — attracted to social impact, pattern of purpose-driven work (OF: Ownership 4/5)')
add_bullet('Honest and self-aware — takes accountability for mistakes (OF: Integrity 4/5)')
add_bullet('Willing to take on secondary Community Building role')
add_bullet('Comfortable working independently with minimal direction')

add_heading_styled('9. Areas for Development', level=2)
add_bullet('Communication precision — sometimes verbose, needs questions re-explained (NF: Communication 3/5)')
add_bullet('HR theoretical frameworks — answers are instinct-based, not methodology-based (NF: Problem Solving 3/5)')
add_bullet('Comprehension agility — needed clarification on questions multiple times (OF: Agility 3/5)')
add_bullet('Attention to detail on policy specifics — initially misstated stipend first-tier amount')

add_heading_styled('10. Overall HR Observation', level=2)
doc.add_paragraph(
    'Vidhi Upadhyay is a learning-stage candidate with genuine HR interest, relevant practical '
    'experience, and strong alignment with the organization\'s mission. She is not a "strong HR" '
    'candidate — her answers lack structured frameworks and theoretical depth — but she demonstrates '
    'the interpersonal instincts, ownership mindset, and adaptability that form a solid foundation '
    'for growth. The internal interviewer assessment ("good, not very good — but learning is a '
    'process that will never end") is an accurate summary. Her full availability, 6-month commitment, '
    'willingness to take on the CB secondary role, and comfort with independent work make her '
    'operationally viable.'
)

add_heading_styled('11. Recommendation', level=2)
p = doc.add_paragraph()
run = p.add_run('Can Be Considered')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
doc.add_paragraph(f'Score: {total_v} / 10')
doc.add_paragraph(
    f'Basis: Strongest evaluated candidate in this batch with full-round data. Natural Fit average '
    f'{nf_avg_v}/5 driven by Empathy (4) and Discretion/Professionalism (4). Organizational Fit '
    f'average {of_avg_v}/5 driven by Ownership (4), Innovation (4), and Integrity (4). '
    f'Areas for development (communication precision, HR frameworks) are addressable through onboarding.'
)

doc.add_paragraph()
add_italic_note(
    'Internal interviewer consensus: Sushma — "good." Nausheen — liked agility and authenticity '
    'answers. Suhani — "okay in natural fit, not 3.5 or 4 level, but good." Decision: Proceed '
    'to check CB fit before onboarding.', 9.5
)

doc.add_paragraph()
doc.add_paragraph()
add_italic_note('Final decision authority: HR OL — Nitin Sir')
add_italic_note('This report is a scoring input, not a hiring decision.')
add_italic_note('To check rejection criteria (H1-H5, S1-S7) for any candidate, use /justify-rejection.')
add_italic_note('Omysha Foundation — Confidential')

# ============================================================
# SAVE
# ============================================================
output_dir = r'd:\workspace\claude-skills\data\evaluation-reports'
output_path = os.path.join(output_dir, 'batch-2026-03-hr-interview-report.docx')
doc.save(output_path)
print(f"Report saved to: {output_path}")
